"""
tools/build_user_guide_doc.py
=============================
Converts USER_GUIDE.md into USER_GUIDE.docx using python-docx.
Run once after editing USER_GUIDE.md:

    python tools/build_user_guide_doc.py
"""

import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("python-docx not installed.  Run: pip install python-docx")
    sys.exit(1)

ROOT   = Path(__file__).parent.parent
MD_IN  = ROOT / "USER_GUIDE.md"
DOC_OUT = ROOT / "USER_GUIDE.docx"

# ─────────────────────────────────────────────────────────────────────────────
# Helpers
# ─────────────────────────────────────────────────────────────────────────────

def _set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def _shade_cell(cell, hex_color="D9E1F2"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _add_code_block(doc, lines):
    """Add a shaded paragraph block for code."""
    for line in lines:
        p = doc.add_paragraph(line)
        p.style = "No Spacing"
        run = p.runs[0] if p.runs else p.add_run(line)
        _set_font(run, name="Courier New", size=9, color=(50, 50, 50))
        # light grey shading on the paragraph
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F2F2F2")
        pPr.append(shd)


def _inline_code(para, text):
    """Add a run styled as inline code."""
    run = para.add_run(text)
    _set_font(run, name="Courier New", size=10, color=(180, 0, 0))
    return run


def _add_para_with_inline(doc, text, base_size=11):
    """Add a paragraph that handles `backtick` inline code spans."""
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    parts = re.split(r"`([^`]+)`", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        if i % 2 == 1:               # inside backticks
            run = para.add_run(part)
            _set_font(run, name="Courier New", size=base_size - 1, color=(180, 0, 0))
        else:
            # strip leading markdown bold/italic markers for body text
            clean = re.sub(r"\*\*(.+?)\*\*", r"\1", part)
            clean = re.sub(r"\*(.+?)\*",   r"\1", clean)
            run = para.add_run(clean)
            _set_font(run, size=base_size)
    return para


# ─────────────────────────────────────────────────────────────────────────────
# Main parser / builder
# ─────────────────────────────────────────────────────────────────────────────

def build_doc(md_path: Path, out_path: Path) -> None:
    doc  = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.1)
        section.right_margin  = Inches(1.1)

    lines      = md_path.read_text(encoding="utf-8").splitlines()
    i          = 0
    in_code    = False
    code_buf: list = []
    in_table   = False
    table_rows: list = []

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        # filter out separator rows (---|---)
        data_rows = [r for r in table_rows if not re.match(r"^\s*\|?[-:| ]+\|?\s*$", r)]
        if not data_rows:
            in_table = False
            table_rows = []
            return

        cols_list = [
            [c.strip() for c in re.split(r"\|", row.strip().strip("|"))]
            for row in data_rows
        ]
        max_cols = max(len(r) for r in cols_list)
        # pad
        for r in cols_list:
            while len(r) < max_cols:
                r.append("")

        tbl = doc.add_table(rows=len(cols_list), cols=max_cols)
        tbl.style = "Table Grid"
        for ri, row_data in enumerate(cols_list):
            for ci, cell_text in enumerate(row_data):
                cell = tbl.rows[ri].cells[ci]
                # strip inline backticks for tables
                clean = re.sub(r"`([^`]+)`", r"\1", cell_text)
                clean = re.sub(r"\*\*(.+?)\*\*", r"\1", clean)
                cell.text = clean
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else None
                if run:
                    _set_font(run, size=9, bold=(ri == 0))
                if ri == 0:
                    _shade_cell(cell, "BDD7EE")
                elif ri % 2 == 0:
                    _shade_cell(cell, "F2F8FF")
        doc.add_paragraph()  # spacing after table
        in_table   = False
        table_rows = []

    while i < len(lines):
        line = lines[i]

        # ── Code block ────────────────────────────────────────────────────────
        if line.strip().startswith("```"):
            if not in_code:
                in_code  = True
                code_buf = []
            else:
                _add_code_block(doc, code_buf)
                doc.add_paragraph()
                in_code  = False
                code_buf = []
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if line.strip().startswith("|"):
            in_table = True
            table_rows.append(line)
            i += 1
            continue
        elif in_table:
            flush_table()

        # ── Headings ──────────────────────────────────────────────────────────
        m = re.match(r"^(#{1,4})\s+(.*)", line)
        if m:
            level = len(m.group(1))
            text  = re.sub(r"`([^`]+)`", r"\1", m.group(2))  # strip backticks
            heading_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            para = doc.add_heading(text, level=level)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^---+\s*$", line):
            doc.add_paragraph("─" * 60).paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        if line.startswith("> "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4)
            para.paragraph_format.space_after = Pt(6)
            run  = para.add_run(line[2:])
            _set_font(run, size=11, italic=True, color=(80, 80, 80))
            i += 1
            continue

        # ── Bullet list ───────────────────────────────────────────────────────
        m = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if m:
            indent = len(m.group(1))
            text   = m.group(2)
            para   = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.left_indent = Inches(0.25 + indent * 0.2)
            para.paragraph_format.space_after = Pt(2)
            # clear auto-added text and add with inline code support
            for run in para.runs:
                run.text = ""
            parts = re.split(r"`([^`]+)`", text)
            for pi, part in enumerate(parts):
                if not part:
                    continue
                if pi % 2 == 1:
                    r = para.add_run(part)
                    _set_font(r, name="Courier New", size=10, color=(180, 0, 0))
                else:
                    clean = re.sub(r"\*\*(.+?)\*\*", r"\1", part)
                    r = para.add_run(clean)
                    _set_font(r, size=11)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Body paragraph ────────────────────────────────────────────────────
        _add_para_with_inline(doc, line)
        i += 1

    if in_table:
        flush_table()

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_doc(MD_IN, DOC_OUT)
